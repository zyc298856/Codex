from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re
import shutil
import sys

sys.path.insert(0, str(Path(__file__).resolve().parents[3] / ".vendor"))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from lxml import etree


ROOT = Path(__file__).resolve().parents[2]
THESIS = ROOT / "paper" / "full_thesis_latest_merged.docx"
ARCHIVE = ROOT / "docs" / "thesis_drafting"
SOURCE_XML = Path(r"C:\Users\Tony\Downloads\applsci-13-04402.xml")


SECTION_NAMES = [
    "1. Introduction",
    "3.3. YOLOv5 Network Model Improvement",
    "4.3. Model Lightweighting Experiment",
    "4.7. Actual Edge Device Deployment Testing",
    "5. Conclusions",
]


TRANSLATIONS = {
    "1. Introduction": [
        "可见光图像由于分辨率高、清晰度好，并且包含容易被人眼理解的丰富视觉细节，因此常被用于目标检测任务。然而，这类图像容易受到天气、光照等外部因素影响，图像质量可能下降，并进一步削弱目标检测精度。红外成像技术正是在这一背景下发挥重要作用。红外成像能够克服部分可见光成像限制，使系统在雾天、烟雾、夜间以及其他恶劣环境下仍能采集图像。通过探测物体释放的热辐射，红外成像与可见光成像相比具有独特优势。在道路场景中，红外图像中的目标通常呈现出条带状视觉特征。不过，红外图像也存在目标轮廓不清晰、目标边缘较模糊、整体纹理信息较弱等问题，这些因素都会给后续目标检测带来困难。",
        "传统红外目标检测技术主要是基于模型的方法，例如模板匹配、阈值分割和 Hausdorff 度量等。随着深度学习的发展，近年来基于卷积神经网络的目标检测技术不断出现。这些方法大体可以分为两阶段算法和单阶段算法。两阶段算法通常以 Faster R-CNN 为代表，单阶段算法则包括 SSD 和 YOLO 等。单阶段算法的设计目标是在算法精度和推理速度之间取得相对平衡。YOLO 系列检测算法是一类单阶段检测器，其目标检测效果优于传统方法，因此在许多应用场景中得到广泛使用。近年来，针对红外道路场景的目标检测也出现了许多改进工作。然而，这些方法往往在以下方面仍存在不足：对于红外道路场景中的条带状目标，需要更加合理地组合条带卷积与传统卷积来提取特征；算法的边界框损失函数需要更加准确地适应红外图像中目标边界的回归；同时，模型本身还需要进一步轻量化，才能更适合实际边缘设备部署。因此，Edge-YOLO 文献正是围绕这些问题展开改进。",
        "上述挑战说明，在边缘设备上开展轻量化红外目标检测具有必要性。为解决这些问题，该文提出了 Edge-YOLO，这是一种面向边缘设备设计的红外目标检测算法，结合了轻量化网络结构和注意力机制。该算法的主要改进可以概括为三个方面。第一，重新设计边界框损失函数，并构建带有幂超参数 α 的损失函数，以增强算法对红外目标边缘的检测能力。第二，设计轻量级内容感知上采样算子，使通道数得到压缩，并尽可能保留边缘区域信息。第三，在 YOLOv5 网络结构基础上进行改进，通过引入条形深度可分离卷积注意力模块增强模型对红外图像中条带状目标特征的提取能力，同时使用改进的 ShuffleBlock 替换 YOLOv5 中的 C3 模块，从而降低模型参数量和计算复杂度。这些改进共同服务于边缘端实时部署这一目标。",
        "该文后续章节安排如下：第二节概述基于神经网络的目标检测相关工作，包括 YOLOv5 以及其他红外目标检测算法；第三节详细介绍该文提出的 Edge-YOLO 算法；第四节给出多组实验结果，用于评价 Edge-YOLO 的检测性能；第五节总结全文的主要贡献。这样的组织方式也体现出轻量化检测模型研究的一般思路，即先分析已有检测算法与红外场景问题，再针对模型结构、损失函数和部署平台进行改进，最后通过实验验证模型在精度、速度和边缘设备适配性方面的综合表现。",
    ],
    "3.3. YOLOv5 Network Model Improvement": [
        "YOLOv5 算法主要面向可见光图像领域设计。由于其模型参数量较大、计算需求较高、模型尺寸也较大，因此更适合部署在 GPU 平台上。然而，该文的研究目标是为边缘嵌入式设备构建轻量化目标检测网络，因此采用轻量级网络结构替代 YOLOv5 中较重的 CSPDarknet 主干网络是合理的选择。ShuffleNetv2 是一种具有代表性的轻量级网络模型，它在速度和精度之间取得了较好的平衡。ShuffleBlock 的主要思想是使用两个分支来减少模型参数量，并通过通道分割和通道重排操作增强不同通道之间的信息交互。与较重的主干网络相比，这种轻量级结构更适合嵌入式端部署，因为边缘设备不仅关注理论计算量，也关注模型实际运行时的内存访问、算子支持和实时响应能力。",
        "为了在不增加明显计算开销的前提下增强红外场景中条带状特征的提取能力，并提高模型对红外图像显著特征的感知能力，该文提出将条形深度可分离卷积注意力模块嵌入 ShuffleBlock 中。该模块首先对输入特征建立快捷分支，并使用 5×5 深度卷积聚合局部信息。随后，模块通过多条分支进行卷积处理，包括一组 1×7 和 7×1 的深度条形卷积、一组 1×11 和 11×1 的深度条形卷积以及一条快捷分支。不同尺度的条形卷积能够增强模型对细长目标和方向性结构的表达能力，这对于红外道路场景中特征边界模糊、目标形态拉长的情况具有实际意义。该设计说明，在面向特定应用场景进行模型轻量化时，不能只简单减少层数或通道数，还需要结合目标形态特点设计更合适的特征提取结构。",
        "在 ShuffleBlock 中嵌入条形深度可分离卷积注意力模块后，文中还在右侧分支加入 SENet 作为通道注意力机制，形成改进后的结构。通常情况下，深度卷积前后会使用 1×1 卷积来融合通道信息，并调整通道数量。然而，原始 ShuffleBlock 在右侧分支的深度卷积层前后都使用 1×1 卷积，这会产生一定冗余。为了减少参数量和计算量，该文对结构进行了进一步调整。由此可见，边缘端检测模型的优化不是单一模块替换，而是围绕特征提取能力、注意力机制、参数冗余和硬件部署约束进行协同设计。对于嵌入式目标检测任务而言，这种思路具有参考价值，因为最终系统既需要保证检测精度，也需要保证模型能够在算力、功耗和内存受限的平台上稳定运行。",
    ],
    "4.3. Model Lightweighting Experiment": [
        "通过将 YOLOv5 的主干特征提取网络替换为该文改进后的 ShuffleBlock，也就是形成图 2 所示的 Edge-YOLO，算法模型整体的参数量、计算量和模型大小都可以得到有效降低。文中随后通过表格比较了模型轻量化改进前后的各项参数。这类实验对于边缘设备部署尤其重要，因为嵌入式系统通常不仅受推理算力限制，还受存储空间、内存带宽和功耗预算限制。如果模型本身过大，即使检测精度较高，也可能难以在低功耗开发板上实现稳定实时运行。",
        "实验结果表明，通过使用改进后的 ShuffleBlock 替换主干网络，与 YOLOv5m 相比，Edge-YOLO 的网络参数量减少了 72.2%，计算量减少了 70.3%，模型大小减少了 71.6%。这说明该文提出的方法具有明显的轻量化效果，有助于减少模型所需的存储资源和计算资源，也更适合部署在边缘嵌入式设备上。对于实际工程项目而言，这一结论提示模型优化不能只看精度指标，还应同时关注参数量、FLOPs、模型文件大小以及板端推理速度等综合因素。",
    ],
    "4.7. Actual Edge Device Deployment Testing": [
        "该文使用瑞芯微 RK3588 嵌入式开发板作为验证平台。RK3588 平台配备四核 A76 与四核 A55 组成的八核 CPU，并集成具有 6 TOPS 计算能力的 NPU。其高算力 NPU 支持 INT4、INT8、INT16 和 FP16 混合计算，能够对网络模型推理过程进行加速。该平台的引入说明，边缘检测算法不仅需要在训练服务器或桌面 GPU 上验证效果，还需要在真实嵌入式硬件上进行部署测试。只有当模型能够完成格式转换、被运行时正确加载，并在板端获得可接受的推理速度时，算法才具备实际工程应用价值。",
        "该文中的算法模型和对比算法模型首先被导出为兼容的 ONNX 格式，然后使用 RKNN-Toolkit2 和 rknpu2 工具转换为 RK3588 平台 NPU 支持的 RKNN 模型。在转换过程中，工具链结合非对称混合量化等推理加速方式，使模型能够在 NPU 上运行。随后，这些模型被用于测试集图像推理，并得到性能对比结果。除了使用 NPU 推理外，文中还对仅使用 ARM CPU 推理的情况进行了比较。这样的实验设计具有较强的工程意义，因为它不仅展示模型本身的检测效果，也说明了从训练框架到 ONNX，再到 RKNN 模型和板端运行时之间的完整部署链路。",
        "从实验结果可以看出，由于模型量化，四种模型在 RK3588 平台上的精度均出现了轻微下降。此外，如果仅使用 ARM CPU 进行推理，YOLO-FIRI 等算法的 FPS 低于 1，也就是说每秒能够推理的图像数量不足一张；该文提出的算法在 CPU 上也只有约 1.1 FPS，难以部署到实际应用场景中。使用 NPU 加速后，可以看到各算法推理速度均得到数十倍提升。该文提出的 Edge-YOLO 达到最高 FPS，比 Faster R-CNN 的推理速度提高约 41.1 倍，比 YOLOv5m 提高约 30.9 倍。尽管与 YOLO-FIRI 相比，Edge-YOLO 的推理速度仅提高约 3.3 倍，但其检测精度更高。因此，文中认为 Edge-YOLO 在保证检测速度的同时也保持了较好的精度表现。",
    ],
    "5. Conclusions": [
        "该文提出的 Edge-YOLO 是一种轻量化红外目标检测方法，目标是在道路场景中保证较好检测性能，同时适合部署在边缘嵌入式设备上。该算法使用改进后的 EX-IoU 边界框损失函数，以提升边界框回归精度；同时采用改进的 CAU-Lite 上采样算子，通过感知周围语义信息来恢复丢失的特征。文中还将条形深度可分离卷积注意力模块嵌入 ShuffleBlock，从而增强模型对红外图像中关键特征的感知能力。实验结果表明，Edge-YOLO 在红外数据集上具有良好的检测效果，并且在参数量、计算量和模型大小方面相较 YOLOv5m 有明显降低。此外，在 RK3588 开发板上进行的边缘部署测试表明，轻量化模型结合 NPU 加速能够显著提升推理速度。总体而言，该文证明了面向应用场景和硬件平台共同设计轻量化检测网络的必要性，也说明边缘目标检测系统需要同时考虑模型结构、特征表达、模型转换、量化影响和实际板端运行效率。",
    ],
}


def extract_sections() -> dict[str, list[str]]:
    root = etree.parse(str(SOURCE_XML)).getroot()
    extracted: dict[str, list[str]] = {}
    for want in SECTION_NAMES:
        for sec in root.xpath('.//*[local-name()="sec"]'):
            title = " ".join("".join(sec.xpath('./*[local-name()="title"]//text()')).split())
            if title == want:
                paras = []
                for p in sec.xpath('./*[local-name()="p"]'):
                    txt = " ".join("".join(p.xpath(".//text()")).split())
                    paras.append(txt)
                extracted[want] = paras
                break
        if want not in extracted:
            raise RuntimeError(f"Section not found: {want}")
    return extracted


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def paragraph_text_from_element(p_elm) -> str:
    texts = []
    for t in p_elm.xpath(".//w:t"):
        if t.text:
            texts.append(t.text)
    return "".join(texts).strip()


def remove_after_until(anchor: Paragraph, end_prefix: str) -> None:
    """Remove all body siblings after anchor until the next paragraph starts with end_prefix."""
    node = anchor._p.getnext()
    while node is not None:
        nxt = node.getnext()
        tag = node.tag
        if tag == qn("w:p") and paragraph_text_from_element(node).startswith(end_prefix):
            break
        node.getparent().remove(node)
        node = nxt


def set_paragraph_font(paragraph: Paragraph, *, english: bool, bold: bool = False, size_pt: float = 12.0) -> None:
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        if english:
            run.font.name = "Times New Roman"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
        else:
            run.font.name = "宋体"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")


def cjk_count(text: str) -> int:
    return len(re.findall(r"[\u4e00-\u9fff]", text))


def main() -> None:
    sections = extract_sections()

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = ARCHIVE / f"full_thesis_latest_merged_before_appendix_a_true_source_{timestamp}.docx"
    ARCHIVE.mkdir(parents=True, exist_ok=True)
    shutil.copy2(THESIS, backup)

    doc = Document(str(THESIS))
    anchor = next(p for p in doc.paragraphs if p.text.strip() == "一、英文原文")

    # Remove current Appendix A body after "一、英文原文" and before Appendix B.
    # This element-level loop is intentionally robust for repeated runs.
    remove_after_until(anchor, "附录B")
    items: list[tuple[str, str, bool, bool]] = []
    # tuple: text, kind, english, bold
    items.append((
        "Source: Li J., Ye J. Edge-YOLO: Lightweight Infrared Object Detection Method Deployed on Edge Devices. Applied Sciences, 2023, 13(7), 4402. https://doi.org/10.3390/app13074402. The following English excerpts are selected from Sections 1, 3.3, 4.3, 4.7, and 5 of the original open-access article.",
        "body", True, False,
    ))
    for sec_name in SECTION_NAMES:
        items.append((sec_name, "sub", True, True))
        for para in sections[sec_name]:
            items.append((para, "body", True, False))

    items.append(("二、英文翻译", "sub", False, True))
    items.append(("Edge-YOLO：部署于边缘设备的轻量化红外目标检测方法（节选翻译）", "sub", False, True))
    for sec_name in SECTION_NAMES:
        cn_title = {
            "1. Introduction": "1. 引言",
            "3.3. YOLOv5 Network Model Improvement": "3.3 YOLOv5 网络模型改进",
            "4.3. Model Lightweighting Experiment": "4.3 模型轻量化实验",
            "4.7. Actual Edge Device Deployment Testing": "4.7 实际边缘设备部署测试",
            "5. Conclusions": "5. 结论",
        }[sec_name]
        items.append((cn_title, "sub", False, True))
        for para in TRANSLATIONS[sec_name]:
            items.append((para, "body", False, False))

    last = anchor
    for text, kind, english, bold in items:
        p = insert_paragraph_after(last, text, style="Normal")
        set_paragraph_font(p, english=english, bold=bold, size_pt=12.0 if kind != "source" else 10.5)
        last = p

    doc.save(str(THESIS))

    translated = "\n".join(p for sec in SECTION_NAMES for p in TRANSLATIONS[sec])
    original = "\n".join(p for sec in SECTION_NAMES for p in sections[sec])
    print(f"backup={backup}")
    print(f"original_words={len(original.split())}")
    print(f"translation_cjk_chars={cjk_count(translated)}")
    print(f"saved={THESIS}")


if __name__ == "__main__":
    main()
