from __future__ import annotations

import re
import sys
from pathlib import Path


ROOT = Path(__file__).resolve().parents[2]
VENDOR = ROOT.parents[1] / ".vendor"
if VENDOR.exists():
    sys.path.insert(0, str(VENDOR))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SOURCE_MD = ROOT / "docs" / "thesis_drafting" / "chapter_3_system_design_draft.md"
OUTPUT_DOCX = ROOT / "docs" / "thesis_drafting" / "chapter_3_system_design_paper_style.docx"
FIGURES_DIR = ROOT / "docs" / "thesis_drafting" / "figures_paper_style"
FIGURE_MAP = {
    "图 3.1": FIGURES_DIR / "fig_3_1_system_architecture_paper.png",
    "图 3.2": FIGURES_DIR / "fig_3_2_model_migration_paper.png",
    "图 3.3": FIGURES_DIR / "fig_3_3_realtime_pipeline_paper.png",
    "图 3.4": FIGURES_DIR / "fig_3_4_multi_context_paper.png",
}


def normalize_text(text: str) -> str:
    text = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[A-Za-z0-9])", "", text)
    text = re.sub(r"(?<=[A-Za-z0-9])\s+(?=[\u4e00-\u9fff])", "", text)
    return text


def set_run_font(run, east_asia: str = "宋体", ascii_font: str = "Times New Roman", size: float = 12, bold: bool = False) -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line: bool = True,
    east_asia: str = "宋体",
    ascii_font: str = "Times New Roman",
    size: float = 12,
    bold: bool = False,
    before: float = 0,
    after: float = 0,
    line_spacing: float = 1.5,
):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing
    if first_line:
        pf.first_line_indent = Pt(24)
    run = p.add_run(normalize_text(text))
    set_run_font(run, east_asia=east_asia, ascii_font=ascii_font, size=size, bold=bold)
    return p


def add_heading(doc: Document, text: str, level: int) -> None:
    if level == 1:
        add_paragraph(
            doc,
            text,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_line=False,
            east_asia="黑体",
            size=18,
            bold=True,
            before=12,
            after=12,
        )
    elif level == 2:
        add_paragraph(
            doc,
            text,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            first_line=False,
            east_asia="黑体",
            size=14,
            bold=True,
            before=12,
            after=6,
        )
    else:
        add_paragraph(
            doc,
            text,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            first_line=False,
            east_asia="黑体",
            size=12,
            bold=True,
            before=9,
            after=4,
        )


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, header: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(normalize_text(text))
    set_run_font(run, east_asia="黑体" if header else "宋体", size=10.5, bold=header)


def set_table_borders(table, color: str = "000000") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single" if edge in ("top", "bottom", "insideH") else "nil")
        element.set(qn("w:sz"), "8" if edge in ("top", "bottom") else "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_caption(doc: Document, text: str, before: float = 6, after: float = 6) -> None:
    add_paragraph(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
        size=10.5,
        before=before,
        after=after,
        line_spacing=1.2,
    )


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.autofit = False
    set_table_borders(table)

    if col_count == 3:
        widths = [Cm(1.4), Cm(3.0), Cm(10.4)]
    elif col_count == 4:
        widths = [Cm(1.3), Cm(2.6), Cm(3.2), Cm(7.7)]
    else:
        widths = [Cm(15 / col_count)] * col_count

    for row_idx, row in enumerate(rows):
        for col_idx in range(col_count):
            cell = table.cell(row_idx, col_idx)
            cell.width = widths[col_idx]
            if row_idx == 0:
                set_cell_shading(cell, "DCEBF4")
            value = row[col_idx] if col_idx < len(row) else ""
            align = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 or len(value) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cell, value, header=row_idx == 0, align=align)
    doc.add_paragraph()


def add_figure(doc: Document, caption_text: str) -> None:
    image_path = None
    for prefix, path in FIGURE_MAP.items():
        if caption_text.startswith(prefix):
            image_path = path
            break

    if not image_path or not image_path.exists():
        add_caption(doc, caption_text)
        add_paragraph(doc, "（图像文件暂未生成）", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(15.0))
    add_caption(doc, caption_text, before=5, after=8)


def is_table_caption(lines: list[str], i: int) -> bool:
    if not re.match(r"^表\s*3\.\d+", lines[i].strip()):
        return False
    j = i + 1
    while j < len(lines) and not lines[j].strip():
        j += 1
    return j < len(lines) and lines[j].strip().startswith("|")


def is_figure_caption(lines: list[str], i: int) -> bool:
    if not re.match(r"^图\s*3\.\d+", lines[i].strip()):
        return False
    j = i + 1
    while j < len(lines) and not lines[j].strip():
        j += 1
    return j < len(lines) and lines[j].strip().startswith("```")


def add_equation(doc: Document, text: str) -> None:
    if "|" in text:
        formula, number = [part.strip() for part in text.rsplit("|", 1)]
    else:
        formula, number = text.strip(), ""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    set_table_borders(table, color="FFFFFF")
    table.cell(0, 0).width = Cm(12.0)
    table.cell(0, 1).width = Cm(2.5)
    set_cell_text(table.cell(0, 0), formula, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 1), number, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def parse_markdown_into_doc(doc: Document, md: str) -> None:
    lines = md.splitlines()
    i = 0
    pending_table_caption: str | None = None
    pending_figure_caption: str | None = None

    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            add_heading(doc, stripped[2:].strip(), 1)
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), 2)
            i += 1
            continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 3)
            i += 1
            continue

        if stripped == "[[PAGEBREAK]]":
            add_page_break(doc)
            i += 1
            continue

        if stripped.startswith("$$"):
            formula_lines: list[str] = []
            first = stripped[2:].strip()
            if first:
                formula_lines.append(first)
            i += 1
            while i < len(lines) and not lines[i].strip().endswith("$$"):
                formula_lines.append(lines[i].strip())
                i += 1
            if i < len(lines):
                last = lines[i].strip()[:-2].strip()
                if last:
                    formula_lines.append(last)
                i += 1
            add_equation(doc, " ".join(formula_lines))
            continue

        if is_table_caption(lines, i):
            pending_table_caption = stripped
            i += 1
            continue
        if is_figure_caption(lines, i):
            pending_figure_caption = stripped
            i += 1
            continue

        if stripped.startswith("```"):
            fence = stripped
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            if "mermaid" in fence and pending_figure_caption:
                add_figure(doc, pending_figure_caption)
                pending_figure_caption = None
            else:
                for code_line in "\n".join(code_lines).strip().splitlines():
                    if code_line.strip():
                        add_equation(doc, code_line.strip())
            continue

        if stripped.startswith("|"):
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    rows.append(cells)
                i += 1
            if pending_table_caption:
                add_caption(doc, pending_table_caption)
                pending_table_caption = None
            add_table(doc, rows)
            continue

        if re.match(r"^\d+\.\s", stripped):
            add_paragraph(doc, stripped)
            i += 1
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt.startswith("#")
                or nxt.startswith("|")
                or nxt.startswith("```")
                or nxt == "[[PAGEBREAK]]"
                or nxt.startswith("$$")
                or is_table_caption(lines, i)
                or is_figure_caption(lines, i)
                or re.match(r"^\d+\.\s", nxt)
            ):
                break
            para_lines.append(nxt)
            i += 1
        add_paragraph(doc, "".join(para_lines))


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def main() -> None:
    md = SOURCE_MD.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    parse_markdown_into_doc(doc, md)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
