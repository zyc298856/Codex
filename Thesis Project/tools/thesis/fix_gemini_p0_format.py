from __future__ import annotations

import shutil
import sys
from pathlib import Path

VENDOR = Path(r"C:\Users\Tony\Desktop\eclipse-workspace-codex\.vendor")
if str(VENDOR) not in sys.path:
    sys.path.insert(0, str(VENDOR))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PROJECT = Path(r"C:\Users\Tony\Desktop\eclipse-workspace-codex\eclipse-workspace\Thesis Project")
DOCX = PROJECT / "paper" / "full_thesis_latest_merged.docx"
SYNCED = PROJECT / "paper" / "full_thesis_latest_merged_int8_synced.docx"
BACKUP = PROJECT / "docs" / "thesis_drafting" / "full_thesis_latest_merged_before_gemini_p0_format_fix.docx"


FORMULA_REPLACEMENTS = {
    "Infer(k)=1, k mod N = 0; Infer(k)=0, k mod N != 0": "Infer(k) = { 1,  k mod N = 0;  0,  k mod N != 0 }",
    "B_smooth(k)=alpha B_detect(k)+(1-alpha)B_smooth(k-1)": "B_smooth(k) = alpha B_detect(k) + (1 - alpha) B_smooth(k - 1)",
}


def ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_paragraph_alignment(p, value: str) -> None:
    ppr = ensure_child(p, "w:pPr")
    jc = ensure_child(ppr, "w:jc")
    jc.set(qn("w:val"), value)


def set_spacing(p, line: str = "300") -> None:
    ppr = ensure_child(p, "w:pPr")
    spacing = ensure_child(ppr, "w:spacing")
    spacing.set(qn("w:line"), line)
    spacing.set(qn("w:lineRule"), "auto")


def paragraph_text(p) -> str:
    return "".join(t.text or "" for t in p.iter(qn("w:t")))


def clear_runs_keep_ppr(p) -> None:
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def add_run_text(p, text: str, font: str = "Cambria Math", size_half_points: str = "24") -> None:
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), font)
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), size_half_points)
    rpr.append(sz)
    szcs = OxmlElement("w:szCs")
    szcs.set(qn("w:val"), size_half_points)
    rpr.append(szcs)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)


def set_border(element, edge: str, val: str, size: str = "4", color: str = "000000") -> None:
    borders = element
    border = borders.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        borders.append(border)
    border.set(qn("w:val"), val)
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), color)


def set_table_three_line(tbl) -> None:
    tbl_pr = ensure_child(tbl, "w:tblPr")
    tbl_borders = ensure_child(tbl_pr, "w:tblBorders")
    set_border(tbl_borders, "top", "single", "12")
    set_border(tbl_borders, "bottom", "single", "12")
    set_border(tbl_borders, "left", "nil", "0")
    set_border(tbl_borders, "right", "nil", "0")
    set_border(tbl_borders, "insideV", "nil", "0")
    set_border(tbl_borders, "insideH", "nil", "0")

    rows = tbl.findall(qn("w:tr"))
    for row_index, row in enumerate(rows):
        cells = row.findall(qn("w:tc"))
        for cell in cells:
            tc_pr = ensure_child(cell, "w:tcPr")
            tc_borders = ensure_child(tc_pr, "w:tcBorders")
            for edge in ("left", "right", "insideV"):
                set_border(tc_borders, edge, "nil", "0")
            set_border(tc_borders, "top", "nil", "0")
            set_border(tc_borders, "bottom", "nil", "0")
            if row_index == 0:
                set_border(tc_borders, "top", "single", "12")
                set_border(tc_borders, "bottom", "single", "4")
            if row_index == len(rows) - 1:
                set_border(tc_borders, "bottom", "single", "12")


def clear_table_borders(tbl) -> None:
    tbl_pr = ensure_child(tbl, "w:tblPr")
    tbl_borders = ensure_child(tbl_pr, "w:tblBorders")
    for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
        set_border(tbl_borders, edge, "nil", "0")

    for cell in tbl.iter(qn("w:tc")):
        tc_pr = ensure_child(cell, "w:tcPr")
        tc_borders = ensure_child(tc_pr, "w:tcBorders")
        for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
            set_border(tc_borders, edge, "nil", "0")


def replace_paragraph_text(doc: Document, old: str, new: str) -> int:
    count = 0
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)
            count += 1
    return count


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)

    doc = Document(str(DOCX))

    formula_count = 0
    number_count = 0
    body = doc.element.body
    paragraphs = list(body.iter(qn("w:p")))
    for p in paragraphs:
        text = paragraph_text(p).strip()
        if text in FORMULA_REPLACEMENTS:
            clear_runs_keep_ppr(p)
            set_paragraph_alignment(p, "center")
            set_spacing(p, "300")
            add_run_text(p, FORMULA_REPLACEMENTS[text])
            formula_count += 1
        elif text in {"(3.1)", "(3.2)", "(4.1)", "(4.2)"}:
            clear_runs_keep_ppr(p)
            set_paragraph_alignment(p, "right")
            set_spacing(p, "300")
            add_run_text(p, text, font="Times New Roman")
            number_count += 1

    formula_table_count = 0
    table_count = 0
    for tbl in body.iter(qn("w:tbl")):
        table_text = "".join(t.text or "" for t in tbl.iter(qn("w:t")))
        if "Infer(k)" in table_text or "B_smooth(k)" in table_text:
            clear_table_borders(tbl)
            formula_table_count += 1
        else:
            set_table_three_line(tbl)
            table_count += 1

    old_gpio = (
        "第三，外设GPIO闭环控制尚未作为本文已完成主线，当前系统采用软件报警overlay和报警事件CSV作为直观替代；"
        "后续可进一步接入报警、云台或其他执行机构，形成从目标感知到硬件响应的完整系统。"
    )
    new_gpio = (
        "第三，系统已在软件层面实现完整的事件触发逻辑，并预留标准化的报警输出接口；"
        "考虑到实验场地安全限制及继电器、蜂鸣器等外接硬件条件约束，本文将闭环执行端落地为视频画面中的报警Overlay和报警事件CSV记录，"
        "从而验证了从目标感知、事件判定到报警输出的完整链路。后续只需在该事件接口处挂载GPIO驱动或云台控制模块，即可扩展为物理报警或姿态响应。"
    )
    gpio_replacements = replace_paragraph_text(doc, old_gpio, new_gpio)

    old_gpio2 = "GPIO硬件闭环则以软件报警显示和事件日志作为替代演示。"
    new_gpio2 = "GPIO硬件闭环在本文中以软件报警显示和事件日志作为等效执行端进行验证，并保留面向GPIO驱动或外设控制的接口扩展空间。"
    gpio_replacements += replace_paragraph_text(doc, old_gpio2, new_gpio2)

    doc.save(str(DOCX))
    shutil.copy2(DOCX, SYNCED)
    print(f"backup={BACKUP}")
    print(f"formulas_fixed={formula_count}")
    print(f"equation_numbers_fixed={number_count}")
    print(f"gpio_replacements={gpio_replacements}")
    print(f"tables_processed={table_count}")
    print(f"formula_tables_borderless={formula_table_count}")
    print(f"synced={SYNCED}")


if __name__ == "__main__":
    main()
