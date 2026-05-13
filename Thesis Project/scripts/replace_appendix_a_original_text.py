from pathlib import Path
import re

from lxml import etree
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"C:\Users\Tony\Desktop\eclipse-workspace-codex\eclipse-workspace\Thesis Project")
DOCX = ROOT / "paper" / "full_thesis_latest_merged.docx"
XML = Path(r"C:\Users\Tony\Downloads\applsci-13-04402.xml")
IMG_DIR = ROOT / "docs" / "thesis_drafting" / "appendix_a_pdfimages"
FORMULA_DIR = ROOT / "docs" / "thesis_drafting" / "edge_yolo_pdf_pages_20260507" / "formula_crops_clean"

FIG_IMAGES = {
    "Figure 1": IMG_DIR / "img-005.png",
    "Figure 2": IMG_DIR / "img-006.png",
    "Figure 3": IMG_DIR / "img-007.png",
    "Figure 4": IMG_DIR / "img-008.png",
    "Figure 5": IMG_DIR / "img-009.png",
    "Figure 6": IMG_DIR / "img-010.png",
    "Figure 7": IMG_DIR / "img-011.png",
    "Figure 8": IMG_DIR / "img-012.png",
    "Figure 9": IMG_DIR / "img-013.png",
}

FORMULA_IMAGES = {
    "(1)": FORMULA_DIR / "formula1.png",
    "(2)": FORMULA_DIR / "formula2.png",
    "(3)": FORMULA_DIR / "formula3.png",
    "(4)": FORMULA_DIR / "formula4.png",
}


def clean_text(text: str) -> str:
    text = " ".join((text or "").split())
    text = text.replace(" ,", ",").replace(" .", ".").replace(" ;", ";").replace(" :", ":")
    text = text.replace("( ", "(").replace(" )", ")")
    text = text.replace(" %", "%")
    return text


def set_run_font(run, size=12, bold=False, italic=False, font="Times New Roman"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def format_paragraph(p, size=12, bold=False, align=None, first_line=True):
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_line:
        pf.first_line_indent = Inches(0.28)
    else:
        pf.first_line_indent = None
    for run in p.runs:
        set_run_font(run, size=size, bold=bold)


def add_para_before(marker, text="", *, size=12, bold=False, align=None, first_line=True):
    p = marker.insert_paragraph_before(text)
    format_paragraph(p, size=size, bold=bold, align=align, first_line=first_line)
    return p


def add_blank_before(marker):
    p = marker.insert_paragraph_before("")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def add_image_before(marker, image_path: Path, *, max_width_in=5.3):
    p = marker.insert_paragraph_before("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.first_line_indent = None
    with Image.open(image_path) as im:
        w, h = im.size
    # Use width-based scaling. The image keeps aspect ratio in Word.
    width = Inches(max_width_in)
    p.add_run().add_picture(str(image_path), width=width)
    return p


def cell_text(cell):
    return clean_text("".join(cell.itertext()))


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        tag = "w:{}".format(edge)
        element = tcBorders.find(qn(tag))
        if edge_data is None:
            if element is not None:
                tcBorders.remove(element)
            continue
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        for key in ["sz", "val", "color", "space"]:
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def clear_cell_borders(cell):
    set_cell_border(cell)


def set_table_font(table, size=8):
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run_font(r, size=size, font="Times New Roman")


def add_table_before(doc, marker, table_wrap):
    table_el = table_wrap.find(".//table")
    if table_el is None:
        return
    rows_xml = table_el.findall(".//tr")
    if not rows_xml:
        return

    # Estimate full column count from spans.
    max_cols = 0
    for tr in rows_xml:
        count = 0
        for c in list(tr):
            if c.tag in ("th", "td"):
                count += int(c.get("colspan", "1"))
        max_cols = max(max_cols, count)

    table = doc.add_table(rows=len(rows_xml), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    occupied = [[False] * max_cols for _ in rows_xml]
    for r_idx, tr in enumerate(rows_xml):
        c_idx = 0
        for c in list(tr):
            if c.tag not in ("th", "td"):
                continue
            while c_idx < max_cols and occupied[r_idx][c_idx]:
                c_idx += 1
            if c_idx >= max_cols:
                break
            rowspan = int(c.get("rowspan", "1"))
            colspan = int(c.get("colspan", "1"))
            target = table.cell(r_idx, c_idx)
            target.text = cell_text(c)
            for rr in range(r_idx, min(r_idx + rowspan, len(rows_xml))):
                for cc in range(c_idx, min(c_idx + colspan, max_cols)):
                    occupied[rr][cc] = True
            if rowspan > 1 or colspan > 1:
                end_r = min(r_idx + rowspan - 1, len(rows_xml) - 1)
                end_c = min(c_idx + colspan - 1, max_cols - 1)
                try:
                    target.merge(table.cell(end_r, end_c))
                except Exception:
                    pass
            c_idx += colspan

    set_table_font(table, size=7 if max_cols >= 8 else 8)

    # Three-line style: no vertical/internal lines, top line, header line, bottom line.
    for row in table.rows:
        for cell in row.cells:
            clear_cell_borders(cell)
    header_rows = len(table_el.findall("./thead/tr")) or 1
    for cell in table.rows[0].cells:
        set_cell_border(cell, top={"val": "single", "sz": "12", "color": "000000"})
    for cell in table.rows[min(header_rows - 1, len(table.rows) - 1)].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": "8", "color": "000000"})
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": "12", "color": "000000"})

    # Move created table from the document end to the current insertion marker.
    marker._p.addprevious(table._tbl)
    add_blank_before(marker)


def flatten_inline(el):
    parts = []
    if el.text:
        parts.append(el.text)
    for child in el:
        if child.tag in ("disp-formula", "list"):
            if child.tail:
                parts.append(child.tail)
            continue
        parts.append(flatten_inline(child))
        if child.tail:
            parts.append(child.tail)
    return "".join(parts)


def paragraph_items(p_el):
    items = []
    if p_el.text:
        items.append(("text", p_el.text))
    for child in p_el:
        if child.tag == "disp-formula":
            label = child.findtext("label")
            items.append(("formula", label))
            if child.tail:
                items.append(("text", child.tail))
        elif child.tag == "list":
            items.append(("list", child))
            if child.tail:
                items.append(("text", child.tail))
        else:
            items.append(("text", flatten_inline(child)))
            if child.tail:
                items.append(("text", child.tail))
    return items


def add_formula_before(marker, label):
    img = FORMULA_IMAGES.get(label)
    if img and img.exists():
        add_image_before(marker, img, max_width_in=5.15)
    else:
        add_para_before(marker, label or "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)


def add_list_before(marker, list_el):
    for li in list_el.findall("./list-item"):
        label = clean_text(li.findtext("label") or "")
        texts = []
        for p in li.findall("./p"):
            texts.append(clean_text(flatten_inline(p)))
        text = clean_text((label + " " + " ".join(texts)).strip())
        if text:
            add_para_before(marker, text, size=12, first_line=True)


def add_figure_before(marker, fig_el):
    label = fig_el.findtext("label")
    caption = clean_text("".join(fig_el.find("caption").itertext())) if fig_el.find("caption") is not None else ""
    img = FIG_IMAGES.get(label)
    if img and img.exists():
        add_image_before(marker, img, max_width_in=5.35)
    add_para_before(marker, f"{label}. {caption}", size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_blank_before(marker)


def add_table_wrap_before(doc, marker, tw_el):
    label = tw_el.findtext("label")
    caption = clean_text("".join(tw_el.find("caption").itertext())) if tw_el.find("caption") is not None else ""
    cap = add_para_before(marker, f"{label}. {caption}", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    # Table 3 is wide enough that starting it at the page bottom leaves only
    # the header row on one page. Keep the caption and full table together.
    if label == "Table 3":
        cap.paragraph_format.page_break_before = True
    cap.paragraph_format.keep_with_next = True
    add_table_before(doc, marker, tw_el)


def refs_in_element(el, ref_type):
    refs = []
    for x in el.findall(".//xref"):
        if x.get("ref-type") == ref_type:
            rid = x.get("rid")
            if rid:
                refs.append(rid)
    return refs


def process_paragraph(doc, marker, p_el, figures, tables, inserted_figs, inserted_tables):
    buffer = []

    def flush():
        nonlocal buffer
        text = clean_text("".join(buffer))
        buffer = []
        if text:
            add_para_before(marker, text, size=12, first_line=True)

    for kind, value in paragraph_items(p_el):
        if kind == "text":
            buffer.append(value)
        elif kind == "formula":
            flush()
            add_formula_before(marker, value)
        elif kind == "list":
            flush()
            add_list_before(marker, value)
    flush()

    for rid in refs_in_element(p_el, "fig"):
        if rid in figures and rid not in inserted_figs:
            add_figure_before(marker, figures[rid])
            inserted_figs.add(rid)
    for rid in refs_in_element(p_el, "table"):
        if rid in tables and rid not in inserted_tables:
            add_table_wrap_before(doc, marker, tables[rid])
            inserted_tables.add(rid)


def process_section(doc, marker, sec, figures, tables, inserted_figs, inserted_tables, level=1):
    title = sec.findtext("title")
    if title and title.strip() == "Figures and Tables":
        return
    if title:
        size = 12 if level > 1 else 12
        p = add_para_before(marker, clean_text(title), size=size, bold=True, first_line=False)
        p.paragraph_format.keep_with_next = True
    for child in sec:
        if child.tag == "title":
            continue
        if child.tag == "p":
            process_paragraph(doc, marker, child, figures, tables, inserted_figs, inserted_tables)
        elif child.tag == "sec":
            process_section(doc, marker, child, figures, tables, inserted_figs, inserted_tables, level + 1)
        elif child.tag == "list":
            add_list_before(marker, child)


def remove_between(doc, start_para, end_para):
    body = doc._body._element
    children = list(body)
    start = children.index(start_para._p)
    end = children.index(end_para._p)
    for el in children[start + 1 : end]:
        body.remove(el)


def main():
    doc = Document(str(DOCX))
    xml_root = etree.parse(str(XML)).getroot()

    paras = doc.paragraphs
    start_para = next(p for p in paras if clean_text(p.text) == "一、英文原文")
    end_para = next(p for p in paras if clean_text(p.text) == "二、英文翻译")
    remove_between(doc, start_para, end_para)

    marker = end_para

    article_meta = xml_root.find("./front/article-meta")
    title = clean_text("".join(article_meta.find("./title-group/article-title").itertext()))
    authors = []
    for contrib in article_meta.findall("./contrib-group/contrib"):
        name = contrib.find("name")
        if name is not None:
            surname = name.findtext("surname") or ""
            given = name.findtext("given-names") or ""
            authors.append(clean_text(f"{given} {surname}"))
    aff = clean_text("".join(article_meta.find("./aff").itertext())) if article_meta.find("./aff") is not None else ""
    abstract = clean_text("".join(article_meta.find("./abstract").itertext()))
    keywords = "; ".join(clean_text("".join(k.itertext())) for k in article_meta.findall("./kwd-group/kwd"))

    add_para_before(
        marker,
        f"Source: Li J., Ye J. {title}. Applied Sciences, 2023, 13(7), 4402. https://doi.org/10.3390/app13074402.",
        size=12,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=False,
    )
    add_blank_before(marker)
    add_para_before(marker, title, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_para_before(marker, ", ".join(authors), size=12, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    if aff:
        add_para_before(marker, aff, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_blank_before(marker)
    add_para_before(marker, "Abstract", size=12, bold=True, first_line=False)
    add_para_before(marker, abstract, size=12, first_line=True)
    add_para_before(marker, f"Keywords: {keywords}", size=12, first_line=True)
    add_blank_before(marker)

    figures = {fig.get("id"): fig for fig in xml_root.findall(".//fig")}
    tables = {tw.get("id"): tw for tw in xml_root.findall(".//table-wrap")}
    inserted_figs = set()
    inserted_tables = set()

    body = xml_root.find("./body")
    for sec in body.findall("./sec"):
        process_section(doc, marker, sec, figures, tables, inserted_figs, inserted_tables, level=1)

    # Insert any remaining figures/tables that were not first-cited in parsed paragraphs.
    for rid, fig in figures.items():
        if rid not in inserted_figs:
            add_figure_before(marker, fig)
    for rid, tw in tables.items():
        if rid not in inserted_tables:
            add_table_wrap_before(doc, marker, tw)

    back = xml_root.find("./back")
    if back is not None:
        for child in back:
            if child.tag == "sec" and clean_text(child.findtext("title") or "") == "Figures and Tables":
                continue
            if child.tag == "ref-list":
                add_para_before(marker, "References", size=12, bold=True, first_line=False)
                for ref in child.findall("./ref"):
                    text = clean_text("".join(ref.itertext()))
                    if text:
                        add_para_before(marker, text, size=10.5, first_line=False)
            elif child.tag == "notes":
                title = clean_text(child.findtext("title") or "")
                text_parts = []
                for p in child.findall("./p"):
                    text_parts.append(clean_text(flatten_inline(p)))
                text = clean_text(" ".join(text_parts))
                if title:
                    add_para_before(marker, title, size=12, bold=True, first_line=False)
                if text:
                    add_para_before(marker, text, size=12, first_line=True)
            elif child.tag == "fn-group":
                text = clean_text("".join(child.itertext()))
                if text:
                    add_para_before(marker, text, size=10.5, first_line=False)

    doc.save(str(DOCX))
    print(DOCX)


if __name__ == "__main__":
    main()
